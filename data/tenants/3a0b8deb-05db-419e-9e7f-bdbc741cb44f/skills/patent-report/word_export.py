"""Export sections to Word."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt

from io_utils import file_artifact, sanitize_filename


def export_word(
    sections: list[dict[str, Any]],
    *,
    out_dir: Path,
    key_prefix: str,
    title: str,
    filename: str,
    charts: list[dict[str, Any]] | None = None,
    summary: str = "",
    interpretation: str = "",
) -> dict[str, Any]:
    doc = Document()
    doc.add_heading(title, level=0)

    if summary:
        doc.add_heading("总结", level=1)
        doc.add_paragraph(summary)
    if interpretation:
        doc.add_heading("解读", level=1)
        doc.add_paragraph(interpretation)

    chart_by_name = {str(c.get("name") or ""): c for c in charts or []}

    for section in sections:
        section_title = str(section.get("title") or "Section")
        doc.add_heading(section_title, level=1)
        rows = section.get("rows") or []
        if rows:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Label"
            hdr[1].text = "Value"
            for item in rows[:50]:
                row = table.add_row().cells
                row[0].text = str(item.get("label") or "")
                value = item.get("value")
                row[1].text = "" if value is None else str(value)
        elif section.get("text"):
            para = doc.add_paragraph(str(section.get("text")))
            for run in para.runs:
                run.font.size = Pt(11)

        chart = chart_by_name.get(section_title)
        if chart and chart.get("path"):
            doc.add_paragraph("")
            doc.add_picture(str(chart["path"]), width=Inches(5.5))

    fname = sanitize_filename(filename, default="patent-report") + ".docx"
    path = out_dir / fname
    doc.save(path)
    key = f"{key_prefix}/{fname}"
    return file_artifact(path, key, fmt="docx")

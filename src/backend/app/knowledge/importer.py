"""Document import - parse files, chunk, embed, store."""

from __future__ import annotations

import uuid
from pathlib import Path

import httpx
from loguru import logger
from sqlalchemy.ext.asyncio import AsyncSession

from app.knowledge.chunk_store import load_document_chunks, replace_document_chunks
from app.knowledge.chunking import chunk_text, chunk_text_with_parents
from app.knowledge.embedder import embedder_for_config
from app.knowledge.milvus_client import milvus_client
from app.knowledge.rag_config import KnowledgeBaseRagSettings, rag_settings_from_kb
from app.models.knowledge import Document, KnowledgeBase


class DocumentImporter:
    """Import documents, chunk them, generate embeddings, and store."""

    def __init__(
        self,
        rag_settings: KnowledgeBaseRagSettings | None = None,
        db: AsyncSession | None = None,
    ) -> None:
        self.rag_settings = rag_settings or KnowledgeBaseRagSettings()
        self.db = db
        self._embedder = embedder_for_config(self.rag_settings.embedding_config())

    async def import_file(
        self,
        kb_id: str,
        user_id: str,
        file_path: Path,
        original_filename: str,
        kb: KnowledgeBase | None = None,
    ) -> Document:
        """Import a file as a knowledge document."""
        if kb is not None:
            self.rag_settings = rag_settings_from_kb(kb)
            self._embedder = embedder_for_config(self.rag_settings.embedding_config())

        content = await self._read_file(file_path)
        source = file_path.suffix.lower().lstrip(".")

        doc = Document(
            id=str(uuid.uuid4()),
            knowledge_base_id=kb_id,
            title=original_filename,
            content=content,
            source=source,
            status="processing",
        )

        children, parents = self._chunk_with_parents(content)
        if not content.strip():
            doc.status = "failed"
            doc.chunk_count = 0
            return doc

        if self.db is not None:
            self.db.add(doc)
            await self.db.flush()

        if self.db is not None:
            await replace_document_chunks(
                self.db,
                document_id=doc.id,
                knowledge_base_id=kb_id,
                chunks=children,
                parents=parents if any(p for p in parents) else None,
            )

        if not milvus_client._connected:
            doc.chunk_count = len(children)
            doc.status = "indexed"
            logger.info(
                f"Document {doc.title} stored ({len(children)} chunks); "
                "Milvus disabled — keyword/hybrid search via DB chunks"
            )
            return doc

        chunk_count = await self.index_document(doc, user_id=user_id, chunks=children)
        doc.chunk_count = chunk_count
        doc.status = "indexed" if chunk_count > 0 else "failed"

        return doc

    async def import_url(
        self, kb_id: str, user_id: str, url: str, kb: KnowledgeBase | None = None
    ) -> Document:
        """Import content from a URL."""
        if kb is not None:
            self.rag_settings = rag_settings_from_kb(kb)
            self._embedder = embedder_for_config(self.rag_settings.embedding_config())

        try:
            async with httpx.AsyncClient(
                timeout=30.0, follow_redirects=True
            ) as client:
                response = await client.get(url)
                response.raise_for_status()
                content = response.text

            doc = Document(
                id=str(uuid.uuid4()),
                knowledge_base_id=kb_id,
                title=url,
                content=content,
                source="url",
                source_url=url,
                status="processing",
            )

            children, parents = self._chunk_with_parents(content)

            if self.db is not None:
                self.db.add(doc)
                await self.db.flush()
                await replace_document_chunks(
                    self.db,
                    document_id=doc.id,
                    knowledge_base_id=kb_id,
                    chunks=children,
                    parents=parents if any(p for p in parents) else None,
                )

            if not milvus_client._connected:
                doc.chunk_count = len(children)
                doc.status = "indexed" if content.strip() else "failed"
                return doc

            chunk_count = await self.index_document(
                doc, user_id=user_id, chunks=children
            )
            doc.chunk_count = chunk_count
            doc.status = "indexed" if chunk_count > 0 else "failed"

            return doc
        except Exception as e:
            logger.error(f"URL import failed for {url}: {e}")
            raise

    def _validate_embedding_dimension(self, dimension: int) -> None:
        if not milvus_client._connected:
            return
        expected = milvus_client.dimension
        if dimension != expected:
            raise ValueError(
                f"Embedding dimension {dimension} does not match Milvus "
                f"collection dimension {expected}. Update EMBEDDING_DIMENSION "
                f"or recreate the Milvus collection."
            )

    async def reindex_document(
        self,
        doc: Document,
        *,
        user_id: str,
        rechunk: bool = True,
    ) -> int:
        """Re-chunk (optional), re-embed, and replace vectors for one document."""
        if not (doc.content or "").strip():
            doc.status = "failed"
            doc.chunk_count = 0
            return 0

        doc.status = "processing"
        if milvus_client._connected:
            await milvus_client.delete_vectors(doc.id)

        if rechunk:
            children, parents = self._chunk_with_parents(doc.content)
            if self.db is not None:
                await replace_document_chunks(
                    self.db,
                    document_id=doc.id,
                    knowledge_base_id=doc.knowledge_base_id,
                    chunks=children,
                    parents=parents if any(p for p in parents) else None,
                )
            pieces = children
        elif self.db is not None:
            pieces = await load_document_chunks(self.db, doc.id)
            if not pieces:
                children, parents = self._chunk_with_parents(doc.content)
                pieces = children
        else:
            pieces = chunk_text(doc.content, self.rag_settings.chunk_config())

        count = await self.index_document(
            doc,
            user_id=user_id,
            chunks=pieces,
            skip_milvus_delete=True,
        )
        doc.status = "indexed" if count > 0 else "failed"
        doc.chunk_count = count
        return count

    async def index_document(
        self,
        doc: Document,
        *,
        user_id: str,
        chunks: list[str] | None = None,
        skip_milvus_delete: bool = False,
    ) -> int:
        """Chunk and index a document's content into Milvus and chunk store."""
        if not milvus_client._connected:
            logger.warning("Milvus not connected, skipping vector indexing")
            if chunks and self.db is not None:
                await replace_document_chunks(
                    self.db,
                    document_id=doc.id,
                    knowledge_base_id=doc.knowledge_base_id,
                    chunks=chunks,
                )
            return len(chunks or [])

        try:
            if chunks is None:
                chunks, _ = self._chunk_with_parents(doc.content)
            pieces = chunks
            if not pieces:
                return 0

            dim = self.rag_settings.embedding_config().resolved_dimension()
            self._validate_embedding_dimension(dim)

            if self.db is not None:
                await replace_document_chunks(
                    self.db,
                    document_id=doc.id,
                    knowledge_base_id=doc.knowledge_base_id,
                    chunks=pieces,
                )

            if milvus_client._connected and not skip_milvus_delete:
                await milvus_client.delete_vectors(doc.id)

            embeddings = await self._embedder.embed_documents(pieces)
            if embeddings and len(embeddings[0]) != milvus_client.dimension:
                self._validate_embedding_dimension(len(embeddings[0]))

            chunk_count = await milvus_client.insert_vectors(
                document_id=doc.id,
                kb_id=doc.knowledge_base_id,
                user_id=user_id,
                chunks=pieces,
                embeddings=embeddings,
            )

            logger.info(f"Indexed document {doc.id}: {chunk_count} chunks")
            return chunk_count
        except Exception as e:
            logger.error(f"Document indexing failed: {e}")
            raise

    async def mark_kb_indexed(self, kb: KnowledgeBase) -> None:
        from app.knowledge.embedding_fingerprint import embedding_fingerprint

        kb.indexed_embedding_key = embedding_fingerprint(kb)

    def _chunk_with_parents(self, content: str) -> tuple[list[str], list[str | None]]:
        """Chunk text, returning (children, parents)."""
        chunk_cfg = self.rag_settings.chunk_config()
        if chunk_cfg.strategy == "semantic":
            pairs = chunk_text_with_parents(content, chunk_cfg, embedder=self._embedder)
        else:
            pairs = chunk_text_with_parents(content, chunk_cfg)
        children = [c for c, _ in pairs]
        parents = [p for _, p in pairs]
        return children, parents

    async def _read_file(self, file_path: Path) -> str:
        """Read file content based on extension."""
        suffix = file_path.suffix.lower()

        if suffix in {".txt", ".md", ".html", ".htm"}:
            return file_path.read_text(encoding="utf-8", errors="replace")
        elif suffix == ".docx":
            try:
                from docx import Document as DocxDocument
            except ImportError as exc:
                raise RuntimeError(
                    "DOCX parsing requires python-docx"
                ) from exc

            document = DocxDocument(file_path)
            parts = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text and paragraph.text.strip()
            ]
            return "\n\n".join(parts)
        elif suffix == ".doc":
            raise ValueError(
                "Legacy .doc files are not supported; please resave as .docx, .pdf, .md, or .txt"
            )
        elif suffix == ".pdf":
            try:
                import pdfplumber

                with pdfplumber.open(file_path) as pdf:
                    pages = [page.extract_text() for page in pdf.pages]
                    return "\n\n".join(p for p in pages if p)
            except ImportError:
                logger.warning(
                    "pdfplumber not installed, returning placeholder"
                )
                return (
                    f"[PDF file: {file_path.name} - requires pdfplumber for parsing]"
                )
        else:
            return file_path.read_text(encoding="utf-8", errors="replace")
